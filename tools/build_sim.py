# -*- coding: utf-8 -*-
"""build_sim.py — MÔ PHỎNG HÀNH TRÌNH TRỰC QUAN (dry-run) cho từng ngày.

Ý tưởng tiết kiệm token: AI chỉ viết MỘT file JSON nhỏ (kịch bản ngày, ~1–2KB).
Script này (KHÔNG dùng token AI) tính toàn bộ: giờ giấc, cước taxi Yandex, vé metro,
quãng đường đi bộ, ngân sách cộng dồn... rồi kết xuất một trang HTML sinh động
(dòng thời gian từng bước + bản đồ lộ trình + tổng kết).

Chạy:  python3 tools/build_sim.py                 # xử lý mọi mo-phong/sim_*.json
       python3 tools/build_sim.py mo-phong/sim_moscow_day1.json
"""
import json, os, sys, glob, math, datetime, re

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
SIMDIR = os.path.join(ROOT, "mo-phong")
FONT = '-apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,"Noto Sans","Helvetica Neue",Arial,sans-serif'

DEFAULT_CFG = {
    "fx": {"rub_usd": 77, "rub_vnd": 337},           # 1 USD≈77₽ ; 1₽≈337₫ (26000/77)
    "taxi": {"base": 200, "per_km": 22, "per_min": 9, "min_fare": 300, "avg_kmh": 24,
             "road_factor": 1.35, "classes": {"economy": 1.0, "comfort": 1.4, "comfort_plus": 1.7, "minivan": 1.8}},
    "metro_ticket": 62, "child_free_under": 7, "walk_kmh": 4.2, "steps_per_km": 1370,
}
ICON = {"taxi": "🚕", "metro": "🚇", "walk": "🚶", "aeroexpress": "🚆", "train": "🚄",
        "boat": "🛳️", "checkin": "🏨", "visit": "📍", "meal": "🍽️", "rest": "☕"}


def hav(a, b):
    R = 6371.0
    p1, p2 = math.radians(a[0]), math.radians(b[0])
    dp = math.radians(b[0] - a[0]); dl = math.radians(b[1] - a[1])
    x = math.sin(dp / 2) ** 2 + math.cos(p1) * math.cos(p2) * math.sin(dl / 2) ** 2
    return 2 * R * math.asin(math.sqrt(x))


def hm(mins):
    mins = int(round(mins)); h = mins // 60
    return f"{h % 24:02d}:{mins % 60:02d}" + (" +1" if h >= 24 else "")


def dur_txt(mins):
    mins = int(round(mins)); h, m = mins // 60, mins % 60
    return (f"{h} giờ " if h else "") + (f"{m} phút" if m or not h else "")


def money(v):
    return f"{int(round(v)):,}".replace(",", ".")


def esc(s):
    import html
    return html.escape(str(s))


def compute(sim):
    cfg = dict(DEFAULT_CFG); cfg.update(sim.get("config", {}))
    tx = dict(DEFAULT_CFG["taxi"]); tx.update(cfg.get("taxi", {}))
    g = sim.get("group", {})
    adults = g.get("adults", 2); ages = g.get("child_ages", []) or []
    children = g.get("children", len(ages))
    child_pay = sum(1 for a in ages if a >= cfg["child_free_under"]) if ages else children
    payers = adults + child_pay
    pts = sim.get("points", {})

    def coord(k): p = pts.get(k, {}); return [p.get("lat"), p.get("lon")]

    clock = 0
    hh, mm = (sim.get("start_time", "09:00").split(":") + ["0"])[:2]
    clock = int(hh) * 60 + int(mm)
    spend = 0.0; walk_km = 0.0
    taxi_total = metro_total = meal_total = ticket_total = 0.0
    taxi_rides = 0
    order = []  # (key,name) for map, in sequence
    out_steps = []

    def add_pt(k):
        if k in pts and (not order or order[-1][0] != k):
            order.append((k, pts[k].get("name", k)))

    for st in sim.get("steps", []):
        t = st.get("t")
        start = clock
        row = {"icon": "•", "title": "", "meta": "", "cost": 0.0, "extra": "", "cls": t}
        if t == "transfer":
            mode = st.get("mode", "taxi")
            frm, to = st.get("from"), st.get("to")
            add_pt(frm);
            if mode in ("taxi", "aeroexpress", "train", "boat"):
                km = st.get("km")
                if km is None and pts.get(frm) and pts.get(to):
                    km = hav(coord(frm), coord(to)) * tx["road_factor"]
                km = km or 0
                if mode == "taxi":
                    tmin = st.get("min") or (km / tx["avg_kmh"] * 60)
                    cls = st.get("class", "comfort"); mult = tx["classes"].get(cls, 1.0)
                    fare = max(tx["min_fare"], tx["base"] + tx["per_km"] * km + tx["per_min"] * tmin) * mult
                    row["cost"] = fare; taxi_total += fare; taxi_rides += 1
                    row["icon"] = ICON["taxi"]
                    row["title"] = st.get("label") or f"Taxi Yandex ({cls.replace('_',' ')}) → {pts.get(to,{}).get('name',to)}"
                    row["meta"] = f"{km:.1f} km · ~{dur_txt(tmin)} · 1 xe cho cả đoàn"
                    row["extra"] = (f"Ước tính cước: gốc {money(tx['base'])} + {money(tx['per_km'])}₽/km×{km:.1f} + "
                                    f"{money(tx['per_min'])}₽/phút×{int(tmin)} → ×{mult} ({cls})")
                    clock += tmin
                else:
                    tmin = st.get("min", 40); perp = st.get("fare", 0); fare = perp * payers
                    row["cost"] = fare; taxi_total += fare
                    row["icon"] = ICON.get(mode, "🚆")
                    row["title"] = st.get("label") or f"{mode.title()} → {pts.get(to,{}).get('name',to)}"
                    row["meta"] = f"~{dur_txt(tmin)}" + (f" · {payers} vé × {money(perp)}₽" if perp else "")
                    clock += tmin
            elif mode == "metro":
                tmin = st.get("min", 35); cost = cfg["metro_ticket"] * payers
                row["cost"] = cost; metro_total += cost; row["icon"] = ICON["metro"]
                row["title"] = f"Metro → {pts.get(to,{}).get('name',to)}"
                row["meta"] = f"~{dur_txt(tmin)} · {payers} vé × {money(cfg['metro_ticket'])}₽" + (f" · {st['stations']}" if st.get("stations") else "")
                clock += tmin
            elif mode == "walk":
                km = st.get("km", 0.5); tmin = km / cfg["walk_kmh"] * 60; walk_km += km
                row["icon"] = ICON["walk"]; row["title"] = f"Đi bộ → {pts.get(to,{}).get('name',to)}"
                row["meta"] = f"{int(km*1000)} m · ~{dur_txt(tmin)}"; clock += tmin
            add_pt(to)
            if st.get("note"): row["extra"] = (row["extra"] + " · " if row["extra"] else "") + st["note"]
        elif t in ("visit", "checkin", "rest"):
            place = st.get("place"); add_pt(place)
            dur = st.get("dur", 60); clock += dur
            nm = pts.get(place, {}).get("name", st.get("name", place))
            row["icon"] = ICON["checkin"] if t == "checkin" else ICON["visit"]
            row["title"] = ({"checkin": "Nhận phòng · ", "rest": "Nghỉ · "}.get(t, "") + nm)
            wk = st.get("walk_km", 0); walk_km += wk
            metas = [f"⏳ {dur_txt(dur)}"]
            if wk: metas.append(f"🚶 đi bộ trong khu ~{wk:.1f} km")
            tk = st.get("tickets")
            if tk:
                c = tk.get("adult", 0) * adults + tk.get("child", 0) * child_pay
                row["cost"] = c; ticket_total += c
                metas.append(f"🎟️ vé {money(c)}₽ ({tk.get('desc','')})")
            row["meta"] = " · ".join(metas)
            acts = st.get("activities", [])
            if acts:
                row["extra"] = "<ul class='acts'>" + "".join(f"<li>{esc(a)}</li>" for a in acts) + "</ul>"
        elif t == "meal":
            dur = st.get("dur", 60); clock += dur
            per_a = st.get("per_adult", 0); per_c = st.get("per_child", 0)
            cost = per_a * adults + per_c * child_pay
            row["cost"] = cost; meal_total += cost; row["icon"] = ICON["meal"]
            row["title"] = f"{st.get('meal','Bữa ăn')} · {esc(st.get('place_name',''))}"
            row["meta"] = f"⏳ {dur_txt(dur)} · {esc(st.get('cuisine',''))} · ~{money(per_a)}₽/người lớn, {money(per_c)}₽/trẻ → {money(cost)}₽"
            dishes = st.get("dishes", [])
            if dishes:
                chips = "".join(f"<span class='dish'>{esc(d.get('n',''))} · {money(d.get('p',0))}₽</span>" for d in dishes)
                row["extra"] = "<div class='dishes'>" + chips + "</div>"
            if st.get("note"): row["extra"] += f"<div class='mnote'>{esc(st['note'])}</div>"
            if st.get("lat"):
                pts.setdefault("_meal%d" % len(out_steps), {"name": st.get("place_name"), "lat": st["lat"], "lon": st["lon"]})
        spend += row["cost"]
        row["start"] = hm(start); row["end"] = hm(clock); row["run"] = spend
        out_steps.append(row)

    total_min = clock - (int(hh) * 60 + int(mm))
    fx = cfg["fx"]
    summary = {
        "adults": adults, "children": children, "child_pay": child_pay, "payers": payers,
        "start": sim.get("start_time"), "end": hm(clock), "total_min": total_min,
        "spend": spend, "usd": spend / fx["rub_usd"], "vnd": spend * fx["rub_vnd"],
        "per_person": spend / max(1, adults + children),
        "walk_km": walk_km, "walk_min": walk_km / cfg["walk_kmh"] * 60, "steps": walk_km * cfg["steps_per_km"],
        "taxi_total": taxi_total, "metro_total": metro_total, "meal_total": meal_total,
        "ticket_total": ticket_total, "taxi_rides": taxi_rides, "cfg": cfg,
    }
    return out_steps, summary, order, pts


def parse_trip_day(name):
    """sim_<trip>_day<N>.json → ('trip', N). Trả (None, None) nếu không khớp."""
    b = os.path.splitext(os.path.basename(name))[0]
    m = re.match(r"^sim_(.+)_day(\d+)$", b)
    if m:
        return m.group(1), int(m.group(2))
    return None, None


def scan_trips(files=None):
    """Quét mọi mo-phong/sim_*.json, gom theo 'trip' (token giữa tên file).
    Trả dict trip -> list[dict(n, base, spend, city, fx)] đã sắp theo N."""
    files = files or sorted(glob.glob(os.path.join(SIMDIR, "sim_*.json")))
    trips = {}
    for jp in files:
        base = os.path.splitext(os.path.basename(jp))[0]
        trip, n = parse_trip_day(jp)
        if trip is None:
            trip, n = base, 1
        sim = json.load(open(jp, encoding="utf-8"))
        _, S, _, _ = compute(sim)
        trips.setdefault(trip, []).append(
            {"n": n, "base": base, "spend": S["spend"],
             "city": sim.get("city", ""), "fx": S["cfg"]["fx"]})
    for t in trips:
        trips[t].sort(key=lambda x: x["n"])
    return trips


def build_nav(base, trips):
    """Dựng dữ liệu thanh điều hướng ngày cho trang có tên file `base`."""
    trip, _ = parse_trip_day(base)
    if trip is None:
        trip = base
    entries = trips.get(trip, [])
    days, prev, nxt = [], None, None
    for i, e in enumerate(entries):
        cur = (e["base"] == base)
        days.append({"n": e["n"], "file": e["base"] + ".html", "cur": cur})
        if cur:
            if i > 0:
                prev = entries[i - 1]["base"] + ".html"
            if i < len(entries) - 1:
                nxt = entries[i + 1]["base"] + ".html"
    label = (entries[0]["city"] if entries else "") or trip
    total = sum(e["spend"] for e in entries)
    return {"prev": prev, "next": nxt, "days": days,
            "label": label, "count": len(entries), "total": total}


def render_nav(nav, pos):
    """HTML thanh điều hướng ngày (pos='top' dính đỉnh, 'bot' ở chân trang)."""
    if not nav or not nav.get("days"):
        return ""
    if nav.get("prev"):
        prev = f'<a class="nv prev" href="{esc(nav["prev"])}">◀ Ngày trước</a>'
    else:
        prev = '<span class="nv prev disabled">◀ Ngày trước</span>'
    if nav.get("next"):
        nxt = f'<a class="nv next" href="{esc(nav["next"])}">Ngày sau ▶</a>'
    else:
        nxt = '<span class="nv next disabled">Ngày sau ▶</span>'
    dots = ""
    for d in nav["days"]:
        cls = "dp cur" if d["cur"] else "dp"
        aria = ' aria-current="page"' if d["cur"] else ""
        dots += f'<a class="{cls}" href="{esc(d["file"])}"{aria}>{d["n"]}</a>'
    ctr = (f'<div class="ctr"><a class="nv idx" href="index.html">Mục lục</a>'
           f'<div class="picker">{dots}</div></div>')
    return f'<nav class="daynav {pos}">{prev}{ctr}{nxt}</nav>'


def render(sim, steps, S, order, pts, nav=None):
    g = sim.get("group", {})
    grp = f"{S['adults']} người lớn" + (f" + {g.get('children',0)} trẻ em" if g.get("children") else "")
    if g.get("child_ages"): grp += f" (bé {', '.join(str(a) for a in g['child_ages'])} tuổi)"
    # timeline
    tl = ""
    for r in steps:
        cost = f"<span class='pill'>{money(r['cost'])} ₽</span>" if r["cost"] else "<span class='pill free'>miễn phí</span>"
        run = f"<span class='run'>Đã chi: {money(r['run'])} ₽</span>"
        tl += f"""<div class="step">
          <div class="tcol"><div class="tm">{r['start']}</div><div class="dot {esc(r['cls'])}">{r['icon']}</div><div class="tm end">{r['end']}</div></div>
          <div class="body"><div class="head"><h3>{esc(r['title'])}</h3><div class="pills">{cost}{run}</div></div>
          <div class="meta">{r['meta']}</div>{('<div class=ex>'+r['extra']+'</div>') if r['extra'] else ''}</div></div>"""
    # map points
    mpts = [{"lat": pts[k]["lat"], "lon": pts[k]["lon"], "name": nm} for k, nm in order if pts.get(k, {}).get("lat")]
    mjson = json.dumps(mpts, ensure_ascii=False)
    cfg = S["cfg"]; tx = cfg["taxi"]
    fx = cfg["fx"]
    navtop = render_nav(nav, "top"); navbot = render_nav(nav, "bot")
    triptot = ""
    if nav and nav.get("days"):
        tt = nav["total"]
        triptot = (f'<div class="triptot">🧭 Tổng chi phí cả hành trình '
                   f'({esc(nav["label"])}, {nav["count"]} ngày): '
                   f'<b>{money(tt)} ₽</b> (~${money(tt / fx["rub_usd"])} · '
                   f'~{money(tt * fx["rub_vnd"] / 1000)}k₫)</div>')
    tar = (f"Taxi: gốc {money(tx['base'])}₽ + {money(tx['per_km'])}₽/km + {money(tx['per_min'])}₽/phút "
           f"(×hệ số hạng xe; road-factor {tx['road_factor']}, tốc độ TB {tx['avg_kmh']} km/h). "
           f"Metro {money(cfg['metro_ticket'])}₽/vé (bé ≥{cfg['child_free_under']} tuổi tính vé). "
           f"Tỷ giá 1$≈{fx['rub_usd']}₽, 1₽≈{fx['rub_vnd']}₫. Quãng đường = đường chim bay × {tx['road_factor']}.")
    return f"""<!DOCTYPE html><html lang="vi"><head><meta charset="UTF-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1.0"/>
<title>Mô phỏng · {esc(sim.get('city',''))} · Ngày {sim.get('day','')}</title>
<link rel="stylesheet" href="../assets/vendor/leaflet.css"/>
<style>
*{{box-sizing:border-box}}body{{font-family:{FONT};margin:0;background:#f6f7fb;color:#1e2733;line-height:1.55}}
a{{color:#152c4e}}
.hd{{background:linear-gradient(135deg,#152c4e,#0f2138);color:#fff;padding:22px 18px}}
.wrap{{max-width:920px;margin:0 auto;padding:0 16px}}
.k{{letter-spacing:.16em;text-transform:uppercase;font-size:11px;color:#c8a24b;font-weight:800}}
.hd h1{{margin:6px 0 4px;font-size:22px}}.hd .sub{{color:#cfd9e8;font-size:13.5px}}
.tot{{display:grid;grid-template-columns:repeat(auto-fit,minmax(150px,1fr));gap:10px;margin-top:14px}}
.tcard{{background:rgba(255,255,255,.10);border:1px solid rgba(255,255,255,.16);border-radius:12px;padding:10px 12px}}
.tcard b{{display:block;font-size:19px;color:#fff}}.tcard span{{font-size:11.5px;color:#b9c6d8}}
#map{{height:280px;margin:0;border-bottom:1px solid #e4e7ee}}
.sec{{padding:18px 0}}
.step{{display:flex;gap:12px;padding:4px 0}}
.tcol{{flex:0 0 54px;display:flex;flex-direction:column;align-items:center;position:relative}}
.tcol:before{{content:"";position:absolute;top:24px;bottom:-8px;width:2px;background:#dfe4ec}}
.tm{{font-size:11px;color:#5c6773;font-weight:700}}.tm.end{{color:#aab3c0;margin-top:2px}}
.dot{{width:34px;height:34px;border-radius:50%;background:#fff;border:2px solid #c8a24b;display:flex;align-items:center;justify-content:center;font-size:16px;margin:2px 0;z-index:1}}
.dot.transfer{{border-color:#2E86DE}}.dot.meal{{border-color:#E67E22}}.dot.visit{{border-color:#1e9e6a}}.dot.checkin{{border-color:#7C5CFC}}
.body{{flex:1;background:#fff;border:1px solid #e4e7ee;border-radius:12px;padding:11px 14px;margin-bottom:10px;box-shadow:0 3px 10px rgba(20,30,50,.05)}}
.head{{display:flex;justify-content:space-between;gap:10px;align-items:flex-start;flex-wrap:wrap}}
.body h3{{margin:0;font-size:15.5px;color:#152c4e}}
.pills{{display:flex;gap:6px;flex-wrap:wrap;align-items:center}}
.pill{{background:#152c4e;color:#fff;border-radius:999px;padding:2px 10px;font-size:12px;font-weight:800;white-space:nowrap}}
.pill.free{{background:#e7f7ef;color:#116644}}
.run{{font-size:11px;color:#8a94a3;white-space:nowrap}}
.meta{{font-size:13px;color:#41506}}.meta{{color:#42505f;margin-top:3px}}
.ex{{margin-top:7px;font-size:13px;color:#374250}}
.acts{{margin:4px 0 0;padding-left:18px}}.acts li{{margin:2px 0}}
.dishes{{display:flex;gap:6px;flex-wrap:wrap;margin-top:4px}}
.dish{{background:#fff4e6;border:1px solid #f0d9b5;color:#8a5a12;border-radius:999px;padding:2px 9px;font-size:12px;font-weight:600}}
.mnote{{margin-top:5px;color:#5c6773;font-size:12.5px;font-style:italic}}
.summary{{background:#fff;border:1px solid #e4e7ee;border-radius:14px;padding:16px;box-shadow:0 6px 18px rgba(20,30,50,.06);margin:6px 0 16px}}
.summary h2{{margin:0 0 10px;font-size:17px;color:#152c4e}}
.brk{{display:grid;grid-template-columns:1fr 1fr;gap:8px 18px;font-size:13.5px}}
.brk .r{{display:flex;justify-content:space-between;border-bottom:1px dashed #e4e7ee;padding:3px 0}}
.brk b{{color:#152c4e}}
.note{{background:#fff8e9;border:1px solid #f0d98a;color:#6b551a;padding:10px 13px;border-radius:10px;font-size:12.5px;margin-top:12px}}
.top{{margin-top:12px}}.top a{{color:#e8edf5;font-weight:700;font-size:13px;text-decoration:none;margin-right:14px}}
footer{{padding:20px;text-align:center;color:#5c6773;font-size:12px}}
.daynav{{background:#152c4e;display:flex;flex-wrap:wrap;align-items:center;justify-content:space-between;gap:8px;padding:8px 12px}}
.daynav.top{{position:sticky;top:0;z-index:600;border-bottom:2px solid #c8a24b;box-shadow:0 3px 12px rgba(15,33,56,.35)}}
.daynav.bot{{border-top:2px solid #c8a24b;margin-top:8px}}
.daynav .nv{{color:#e8edf5;text-decoration:none;font-weight:700;font-size:13.5px;min-height:40px;display:inline-flex;align-items:center;padding:0 12px;border-radius:9px}}
.daynav a.nv:hover{{background:rgba(255,255,255,.12)}}
.daynav .nv.disabled{{opacity:.32;pointer-events:none}}
.daynav .nv.idx{{color:#c8a24b;border:1px solid rgba(200,162,75,.45)}}
.daynav .ctr{{display:flex;flex-wrap:wrap;align-items:center;justify-content:center;gap:8px;flex:1 1 auto}}
.daynav .picker{{display:flex;flex-wrap:wrap;gap:5px;justify-content:center}}
.daynav .picker .dp{{min-width:40px;min-height:40px;padding:0 8px;display:inline-flex;align-items:center;justify-content:center;background:rgba(255,255,255,.09);color:#e8edf5;text-decoration:none;font-size:13.5px;font-weight:700;border-radius:9px}}
.daynav .picker .dp:hover{{background:rgba(200,162,75,.4)}}
.daynav .picker .dp.cur{{background:#c8a24b;color:#152c4e;font-weight:800}}
.triptot{{margin-top:12px;padding:11px 13px;background:#f4f7fc;border:1px solid #d7e0ee;border-left:4px solid #c8a24b;border-radius:10px;font-size:13.5px;color:#1e2733}}
.triptot b{{color:#152c4e;font-size:15px}}
@media(max-width:560px){{.brk{{grid-template-columns:1fr}}.daynav{{gap:6px;padding:7px 9px}}.daynav .nv{{font-size:12.5px;padding:0 9px}}}}
</style></head><body>
{navtop}
<div class="hd"><div class="wrap">
  <div class="k">Cẩm nang Du lịch Nga · Mô phỏng hành trình (dry-run)</div>
  <h1>{esc(sim.get('city',''))} · Ngày {sim.get('day','')} — {esc(grp)}</h1>
  <div class="sub">{esc(sim.get('date',''))} · bắt đầu {S['start']} → kết thúc {S['end']} · mô phỏng như một đoàn khách thật</div>
  <div class="tot">
    <div class="tcard"><b>{dur_txt(S['total_min'])}</b><span>Tổng thời gian trong ngày</span></div>
    <div class="tcard"><b>{money(S['spend'])} ₽</b><span>≈ ${money(S['usd'])} · ≈ {money(S['vnd']/1000)}k₫ · cả đoàn</span></div>
    <div class="tcard"><b>{S['walk_km']:.1f} km</b><span>Đi bộ (~{int(S['steps'])} bước · {dur_txt(S['walk_min'])})</span></div>
    <div class="tcard"><b>{S['taxi_rides']} cuốc xe</b><span>+ metro · di chuyển trong ngày</span></div>
  </div>
  <div class="top"><a href="../trung-tam.html">🏠 Trang chủ</a><a href="../gis.html">🗺️ Bản đồ GIS</a></div>
</div></div>
<div id="map"></div>
<div class="wrap">
  <div class="sec">{tl}</div>
  <div class="summary"><h2>📊 Tổng kết ngày</h2><div class="brk">
    <div class="r"><span>⏱️ Thời gian</span><b>{dur_txt(S['total_min'])} ({S['start']}–{S['end']})</b></div>
    <div class="r"><span>🚶 Đi bộ</span><b>{S['walk_km']:.1f} km · ~{int(S['steps'])} bước</b></div>
    <div class="r"><span>🚕 Taxi (tổng)</span><b>{money(S['taxi_total'])} ₽ · {S['taxi_rides']} cuốc</b></div>
    <div class="r"><span>🚇 Metro</span><b>{money(S['metro_total'])} ₽</b></div>
    <div class="r"><span>🍽️ Ăn uống</span><b>{money(S['meal_total'])} ₽</b></div>
    <div class="r"><span>🎟️ Vé tham quan</span><b>{money(S['ticket_total'])} ₽</b></div>
    <div class="r"><span>💰 TỔNG CHI (cả đoàn)</span><b>{money(S['spend'])} ₽ ≈ ${money(S['usd'])}</b></div>
    <div class="r"><span>👤 Bình quân/người</span><b>{money(S['per_person'])} ₽ ≈ {money(S['per_person']*fx['rub_vnd']/1000)}k₫</b></div>
  </div>
  {triptot}
  <div class="note">📌 <b>Cơ sở tính (điều chỉnh trong file JSON):</b> {esc(tar)}<br/>Đây là mô phỏng ước tính để hình dung &amp; lên ngân sách — giá thực tế (taxi Yandex theo giờ cao điểm, thực đơn) có thể dao động.</div>
  </div>
</div>
{navbot}
<footer>Cẩm nang Du lịch Nga · Mô phỏng hành trình · dữ liệu do người biên soạn nhập, AI gợi ý quán ăn &amp; tính toán.</footer>
<script src="../assets/vendor/leaflet.js"></script>
<script>
var P={mjson};
var map=L.map('map',{{scrollWheelZoom:false}});
L.tileLayer('https://{{s}}.tile.openstreetmap.org/{{z}}/{{x}}/{{y}}.png',{{maxZoom:18,attribution:'© OpenStreetMap'}}).addTo(map);
var latlngs=[];
P.forEach(function(p,i){{
  var m=L.marker([p.lat,p.lon]).addTo(map).bindPopup((i+1)+'. '+p.name);
  var ic=L.divIcon({{className:'',html:'<div style="background:#152c4e;color:#fff;width:24px;height:24px;border-radius:50%;display:flex;align-items:center;justify-content:center;font:700 12px sans-serif;border:2px solid #fff;box-shadow:0 1px 4px rgba(0,0,0,.4)">'+(i+1)+'</div>',iconSize:[24,24],iconAnchor:[12,12]}});
  m.setIcon(ic); latlngs.push([p.lat,p.lon]);
}});
if(latlngs.length){{ L.polyline(latlngs,{{color:'#c8a24b',weight:3,dashArray:'6,6'}}).addTo(map); map.fitBounds(L.latLngBounds(latlngs).pad(0.2)); }}
</script>
</body></html>"""


def _plain(s):
    return re.sub("<[^>]+>", "", s or "").strip()


def _extra_lines(s):
    if not s: return []
    lines = re.findall(r"<li>(.*?)</li>", s, re.S)
    dishes = re.findall(r"<span class='dish'>(.*?)</span>", s, re.S)
    if dishes: lines.append("Món gợi ý: " + ", ".join(_plain(d) for d in dishes))
    lines += re.findall(r"<div class='mnote'>(.*?)</div>", s, re.S)
    if not lines:
        t = _plain(s)
        if t: lines = [t]
    return [_plain(x) for x in lines if _plain(x)]


def build_sim_docx(sim, steps, S, path):
    """Bản Word để người dùng tự chỉnh sửa mô phỏng."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    NAVY = RGBColor(0x15, 0x2c, 0x4e); GOLD = RGBColor(0xa9, 0x86, 0x3a); MUT = RGBColor(0x5c, 0x67, 0x73)
    g = sim.get("group", {})
    grp = f"{S['adults']} người lớn" + (f" + {g.get('children',0)} trẻ em" if g.get("children") else "")
    if g.get("child_ages"): grp += f" (bé {', '.join(str(a) for a in g['child_ages'])} tuổi)"
    doc = Document(); st = doc.styles["Normal"]; st.font.name = "Arial"; st.font.size = Pt(11)
    p = doc.add_paragraph(); r = p.add_run("CẨM NANG DU LỊCH NGA · MÔ PHỎNG HÀNH TRÌNH (bản Word để chỉnh sửa)")
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = GOLD
    p = doc.add_paragraph(); r = p.add_run(f"{sim.get('city','')} · Ngày {sim.get('day','')} — {grp}")
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = NAVY
    p = doc.add_paragraph(); r = p.add_run(f"{sim.get('date','')} · {S['start']}–{S['end']} · Tổng {money(S['spend'])}₽ (~${money(S['usd'])}) · Đi bộ {S['walk_km']:.1f} km")
    r.italic = True; r.font.color.rgb = MUT
    for r0 in steps:
        h = doc.add_paragraph(); run = h.add_run(f"{r0['start']}–{r0['end']}   {_plain(r0['title'])}")
        run.bold = True; run.font.size = Pt(12.5); run.font.color.rgb = NAVY
        c = doc.add_paragraph(); cc = c.add_run(("Chi phí: " + money(r0["cost"]) + "₽" if r0["cost"] else "Miễn phí") + f"    ·    Đã chi (cộng dồn): {money(r0['run'])}₽")
        cc.font.size = Pt(10); cc.font.color.rgb = GOLD
        if r0["meta"]:
            m = doc.add_paragraph(_plain(r0["meta"])); m.runs[0].font.size = Pt(10.5)
        for line in _extra_lines(r0["extra"]):
            doc.add_paragraph(line, style="List Bullet")
    hd = doc.add_heading(level=1); rr = hd.add_run("Tổng kết ngày"); rr.font.color.rgb = NAVY
    rows = [("Thời gian", f"{dur_txt(S['total_min'])} ({S['start']}–{S['end']})"),
            ("Đi bộ", f"{S['walk_km']:.1f} km (~{int(S['steps'])} bước)"),
            ("Taxi", f"{money(S['taxi_total'])}₽ · {S['taxi_rides']} cuốc"),
            ("Metro", f"{money(S['metro_total'])}₽"), ("Ăn uống", f"{money(S['meal_total'])}₽"),
            ("Vé tham quan", f"{money(S['ticket_total'])}₽"),
            ("TỔNG CHI (cả đoàn)", f"{money(S['spend'])}₽ ≈ ${money(S['usd'])} ≈ {money(S['vnd']/1000)}k₫"),
            ("Bình quân/người", f"{money(S['per_person'])}₽")]
    for label, val in rows:
        p = doc.add_paragraph(); a = p.add_run(f"{label}: "); a.bold = True; a.font.color.rgb = NAVY; p.add_run(val)
    n = doc.add_paragraph(); nr = n.add_run("Đây là mô phỏng ước tính — bạn có thể sửa trực tiếp file Word này, hoặc sửa file JSON (giá taxi/metro/món, thời lượng) rồi chạy lại build_sim.py.")
    nr.italic = True; nr.font.size = Pt(9); nr.font.color.rgb = MUT
    doc.save(path)


def build_sim_index(rows, trips_summary=None):
    e = esc
    trips_summary = trips_summary or []
    tcards = ""; g_total = g_usd = g_vnd = 0.0; g_days = 0
    for t in trips_summary:
        g_total += t["total"]; g_usd += t["usd"]; g_vnd += t["vnd"]; g_days += t["days"]
        tcards += (f'<div class="tsum"><div class="d">{t["days"]} ngày</div>'
                   f'<h3>{e(t["label"])}</h3>'
                   f'<div class="big">{money(t["total"])} ₽</div>'
                   f'<div class="sub">≈ ${money(t["usd"])} · ≈ {money(t["vnd"] / 1000)}k₫</div></div>')
    grand = ""
    if trips_summary:
        grand = (f'<div class="grand"><span>🧭 Tổng chi phí cả {len(trips_summary)} hành trình · '
                 f'{g_days} ngày</span><b>{money(g_total)} ₽ ≈ ${money(g_usd)} ≈ {money(g_vnd / 1000)}k₫</b></div>')
    trips_block = (f'<h2 class="sh">💰 Tổng chi phí theo hành trình</h2>'
                   f'<div class="trips">{tcards}{grand}</div>') if trips_summary else ""
    cards = ""
    for r in sorted(rows, key=lambda x: (str(x["city"]), x["day"])):
        cards += (f'<div class="card"><a class="ml" href="{e(r["file"])}"><div class="d">NGÀY {r["day"]}</div>'
                  f'<h3>{e(r["city"])}</h3><p>{e(r["group"])}</p>'
                  f'<div class="m">⏱ {e(r["dur"])} · 💰 {e(r["spend"])} · 🚶 {e(r["walk"])}</div></a>'
                  f'<div class="dl">📝 Sửa trong Word: <a href="{e(r["docx"])}">{e(r["docx"])}</a></div></div>')
    html = f"""<!DOCTYPE html><html lang="vi"><head><meta charset="UTF-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1.0"/>
<title>Mô phỏng hành trình — Mục lục</title>
<style>*{{box-sizing:border-box}}body{{font-family:{FONT};margin:0;background:#f6f7fb;color:#1e2733}}
.hd{{background:linear-gradient(135deg,#152c4e,#0f2138);color:#fff;padding:24px 18px}}
.hd .k{{letter-spacing:.16em;text-transform:uppercase;font-size:11px;color:#c8a24b;font-weight:800}}
.hd h1{{margin:6px 0 4px;font-size:22px}}.hd a{{color:#e8edf5;font-weight:700;text-decoration:none;font-size:13px}}
.wrap{{max-width:960px;margin:0 auto;padding:18px}}
.grid{{display:grid;grid-template-columns:repeat(auto-fill,minmax(240px,1fr));gap:14px}}
.card{{background:#fff;border:1px solid #e4e7ee;border-radius:14px;padding:15px;text-decoration:none;color:#1e2733;box-shadow:0 5px 16px rgba(20,30,50,.06)}}
.card:hover{{border-color:#152c4e}}.card .ml{{text-decoration:none;color:inherit;display:block}}
.card .d{{font-size:11px;font-weight:800;color:#a9863a;letter-spacing:.08em}}
.card h3{{margin:3px 0 4px;font-size:17px;color:#152c4e}}.card p{{margin:0;font-size:12.5px;color:#5c6773}}
.card .m{{margin-top:8px;font-size:12.5px;color:#374250;font-weight:600}}
.card .dl{{margin-top:6px;font-size:12px;color:#a9863a}}.card .dl a{{color:#a9863a;font-weight:700}}
.sh{{margin:14px 0 10px;font-size:16px;color:#152c4e}}
.trips{{display:grid;grid-template-columns:repeat(auto-fit,minmax(230px,1fr));gap:12px;margin:0 0 20px}}
.tsum{{background:#fff;border:1px solid #e4e7ee;border-left:4px solid #c8a24b;border-radius:12px;padding:14px 16px;box-shadow:0 5px 16px rgba(20,30,50,.06)}}
.tsum .d{{font-size:11px;font-weight:800;color:#a9863a;letter-spacing:.06em;text-transform:uppercase}}
.tsum h3{{margin:3px 0 6px;font-size:17px;color:#152c4e}}
.tsum .big{{font-size:20px;font-weight:800;color:#152c4e}}
.tsum .sub{{font-size:12.5px;color:#5c6773;margin-top:2px}}
.grand{{grid-column:1/-1;background:linear-gradient(135deg,#152c4e,#0f2138);color:#fff;border-radius:12px;padding:14px 16px;display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px}}
.grand b{{color:#c8a24b;font-size:19px}}
.foot{{color:#5c6773;font-size:12px;text-align:center;padding:20px}}</style></head><body>
<div class="hd"><div class="wrap" style="padding:0"><div class="k">Cẩm nang Du lịch Nga · Mô phỏng hành trình</div>
<h1>🎬 Mô phỏng các ngày (dry-run)</h1><a href="../trung-tam.html">🏠 Trang chủ</a></div></div>
<div class="wrap">{trips_block}<div class="grid">{cards}</div>
<div class="foot">Bấm một ngày để xem mô phỏng trực quan; tải .docx để chỉnh sửa trong Word.</div></div></body></html>"""
    with open(os.path.join(SIMDIR, "index.html"), "w", encoding="utf-8") as f:
        f.write(html)


def main():
    args = [a for a in sys.argv[1:] if not a.startswith("-")]
    files = args or sorted(glob.glob(os.path.join(SIMDIR, "sim_*.json")))
    if not files:
        print("Không có file mo-phong/sim_*.json"); return
    os.makedirs(SIMDIR, exist_ok=True)
    trips = scan_trips()  # quét TẤT CẢ sim để dựng điều hướng + tổng cả hành trình
    for fp in files:
        if not os.path.isabs(fp): fp = os.path.join(ROOT, fp)
        sim = json.load(open(fp, encoding="utf-8"))
        steps, S, order, pts = compute(sim)
        base = os.path.splitext(fp)[0]
        nav = build_nav(os.path.basename(base), trips)
        open(base + ".html", "w", encoding="utf-8").write(render(sim, steps, S, order, pts, nav))
        try:
            build_sim_docx(sim, steps, S, base + ".docx")
        except Exception as e:
            print("  (docx lỗi:", e, ")")
        print(f"+ {os.path.basename(base)}.html/.docx · {money(S['spend'])}₽ · {dur_txt(S['total_min'])} · đi bộ {S['walk_km']:.1f}km")
    # mục lục tất cả mô phỏng
    rows = []
    for jp in sorted(glob.glob(os.path.join(SIMDIR, "sim_*.json"))):
        sim = json.load(open(jp, encoding="utf-8")); steps, S, order, pts = compute(sim)
        g = sim.get("group", {})
        grp = f"{S['adults']} người lớn" + (f" + {g.get('children',0)} trẻ" if g.get("children") else "")
        b = os.path.splitext(os.path.basename(jp))[0]
        rows.append({"file": b + ".html", "docx": b + ".docx", "city": sim.get("city", ""), "day": sim.get("day", ""),
                     "spend": money(S["spend"]) + "₽", "dur": dur_txt(S["total_min"]), "walk": f"{S['walk_km']:.1f}km", "group": grp})
    trips_summary = []
    for trip in sorted(trips):
        entries = trips[trip]
        total = sum(x["spend"] for x in entries)
        fx = entries[0]["fx"]
        trips_summary.append({"trip": trip, "label": (entries[0]["city"] or trip),
                              "days": len(entries), "total": total,
                              "usd": total / fx["rub_usd"], "vnd": total * fx["rub_vnd"]})
    build_sim_index(rows, trips_summary)
    print(f"+ mo-phong/index.html ({len(rows)} ngày)")
    for t in trips_summary:
        print(f"  ∑ {t['trip']} ({t['label']}, {t['days']} ngày): "
              f"{money(t['total'])}₽ ≈ ${money(t['usd'])} ≈ {money(t['vnd']/1000)}k₫")
    print(f"  ∑ TỔNG CỘNG tất cả hành trình: {money(sum(t['total'] for t in trips_summary))}₽")


if __name__ == "__main__":
    main()
