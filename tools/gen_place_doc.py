# -*- coding: utf-8 -*-
"""gen_place_doc.py — Kết xuất TÀI LIỆU THUYẾT MINH NỘI BỘ (Word + HTML) cho từng địa điểm.

Đầu vào: _incoming/doc_<slug>.json (agent soạn — nội dung nguyên gốc + nguồn).
Đầu ra:  tai-lieu-noi-bo/<region>/<slug>.docx  và  <slug>.html
Kèm: nút ĐỌC THÀNH TIẾNG (TTS) trong bản HTML, link quay về trang index & hub.

Chạy: python3 tools/gen_place_doc.py
"""
import json, os, glob, html as _html, datetime, unicodedata

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
IN = os.path.join(ROOT, "_incoming")
REGIONS = os.path.join(ROOT, "data", "regions")
OUTROOT = os.path.join(ROOT, "tai-lieu-noi-bo")
try:
    TODAY = datetime.date.today().isoformat()
except Exception:
    TODAY = "2026-07-24"

# Font tiếng Việt hiển thị đúng (giống app chính) — TRÁNH serif Georgia gây lỗi dấu.
FONT = '-apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,"Noto Sans","Helvetica Neue",Arial,sans-serif'

SLUG2REGION = {}
for f in glob.glob(os.path.join(REGIONS, "*.json")):
    reg = os.path.basename(f)[:-5]
    for r in json.load(open(f, encoding="utf-8")):
        SLUG2REGION[r["slug"]] = reg


def nfc(s):
    return unicodedata.normalize("NFC", s) if isinstance(s, str) else s


def normd(o):
    if isinstance(o, str):
        return nfc(o)
    if isinstance(o, list):
        return [normd(x) for x in o]
    if isinstance(o, dict):
        return {k: normd(v) for k, v in o.items()}
    return o


def _fix_lists(d):
    """Chống lỗi: chuẩn hoá references/images về dạng dict, kể cả khi agent ghi nhầm thành chuỗi."""
    d = dict(d)
    refs = []
    for r in (d.get("references") or []):
        if isinstance(r, dict):
            refs.append(r)
        elif isinstance(r, str):
            refs.append({"title": r, "url": r if r.startswith("http") else ""})
    d["references"] = refs
    imgs = []
    for im in (d.get("images") or []):
        if isinstance(im, dict):
            imgs.append(im)
        elif isinstance(im, str):
            imgs.append({"url": im, "caption": ""})
    d["images"] = imgs
    return d


def region_of(d):
    return SLUG2REGION.get(d.get("slug"), "khac")


TTS_JS = """
<div class="ttsbar no-print">
  <button id="ttsPlay" class="tb">🔊 Nghe toàn bài</button>
  <button id="ttsPause" class="tb" style="display:none">⏸ Tạm dừng</button>
  <button id="ttsStop" class="tb">⏹</button>
  <span class="trate">Tốc độ <input type="range" id="ttsRate" min="0.6" max="1.4" step="0.1" value="1"/></span>
  <span id="ttsWarn" style="display:none;color:#b26a00;font-size:12px">máy chưa có giọng vi-VN — vẫn đọc được</span>
</div>
<script>
(function(){
  var synth = window.speechSynthesis;
  var bar = document.querySelector('.ttsbar');
  if(!synth){ if(bar) bar.style.display='none'; return; }
  var rate=1, chunks=[], idx=0, playing=false, paused=false;
  function collect(){
    var out=[];
    document.querySelectorAll('.doc h1,.doc h2,.doc h3,.doc p,.doc li,.doc figcaption').forEach(function(el){
      var t=(el.innerText||'').trim(); if(!t) return;
      t.split(/(?<=[.!?…:;])\\s+/).forEach(function(s){
        s=s.trim(); if(!s) return;
        while(s.length>200){ out.push(s.slice(0,200)); s=s.slice(200); }
        if(s) out.push(s);
      });
    });
    return out;
  }
  function pickVoice(){
    var vs=synth.getVoices()||[];
    return vs.find(function(v){return /vi[-_]?VN/i.test(v.lang)||/vietnam/i.test(v.name);})
        || vs.find(function(v){return /^vi/i.test(v.lang);}) || null;
  }
  function speakNext(){
    if(!playing) return;
    if(idx>=chunks.length){ stop(); return; }
    var u=new SpeechSynthesisUtterance(chunks[idx]);
    u.lang='vi-VN'; u.rate=rate; var v=pickVoice(); if(v) u.voice=v;
    u.onend=function(){ if(playing&&!paused){ idx++; speakNext(); } };
    synth.speak(u);
  }
  function setBtn(on){ document.getElementById('ttsPlay').style.display=on?'none':''; document.getElementById('ttsPause').style.display=on?'':'none'; }
  function start(){
    if(paused){ paused=false; try{synth.resume();}catch(e){} setBtn(true); return; }
    chunks=collect(); idx=0; playing=true; paused=false; setBtn(true);
    if(!pickVoice()) document.getElementById('ttsWarn').style.display='';
    speakNext();
  }
  function pause(){ if(!playing) return; paused=true; try{synth.pause();}catch(e){} setBtn(false); }
  function stop(){ playing=false; paused=false; try{synth.cancel();}catch(e){} setBtn(false); }
  document.getElementById('ttsPlay').onclick=start;
  document.getElementById('ttsPause').onclick=pause;
  document.getElementById('ttsStop').onclick=stop;
  document.getElementById('ttsRate').oninput=function(){ rate=parseFloat(this.value)||1; };
  if(typeof synth.onvoiceschanged!=='undefined'){ synth.onvoiceschanged=function(){}; }
})();
</script>
"""


def build_html(d):
    d = _fix_lists(normd(d))
    esc = _html.escape
    secs = d.get("sections", [])
    toc = "".join(f'<li><a href="#s{i}">{esc(s.get("heading",""))}</a></li>' for i, s in enumerate(secs))
    body = ""
    for i, s in enumerate(secs):
        paras = "".join(f"<p>{esc(p)}</p>" for p in s.get("paras", []))
        body += f'<section id="s{i}"><h2>{esc(s.get("heading",""))}</h2>{paras}</section>'
    hl = "".join(f"<li>{esc(x)}</li>" for x in d.get("highlights", []))
    hlbox = f'<div class="hl"><h3>★ Điểm nhấn</h3><ul>{hl}</ul></div>' if hl else ""
    imgs = ""
    for im in d.get("images", []):
        u = esc(im.get("url", "")); cap = esc(im.get("caption", ""))
        imgs += f'<figure><img loading="lazy" src="{u}" alt="{cap}"/><figcaption>{cap}</figcaption></figure>'
    imgbox = f'<section><h2>Hình ảnh</h2><div class="gal">{imgs}</div></section>' if imgs else ""
    refs = "".join(f'<li><a href="{esc(r.get("url",""))}" target="_blank" rel="noopener">{esc(r.get("title",""))}</a></li>' for r in d.get("references", []))
    refbox = f'<section><h2>Nguồn tham khảo & liên kết</h2><ul class="refs">{refs}</ul></section>' if refs else ""
    src = "".join(f"<li>{esc(x)}</li>" for x in d.get("sources", []))
    srcbox = f'<section class="src"><h3>Nguồn dữ kiện</h3><ul>{src}</ul></section>' if src else ""
    meta = " · ".join([x for x in [d.get("name_ru"), d.get("name_en")] if x])
    return f"""<!DOCTYPE html><html lang="vi"><head><meta charset="UTF-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1.0"/>
<title>{esc(d.get('name_vi',''))} — Tài liệu thuyết minh nội bộ</title>
<style>
:root{{--navy:#152c4e;--gold:#a9863a;--ink:#22303f;--muted:#5c6773;--line:#e4e7ee}}
*{{box-sizing:border-box}}
body{{font-family:{FONT};color:var(--ink);line-height:1.75;margin:0;background:#fbfcfe;padding-bottom:70px}}
.topbar{{background:#152c4e;color:#fff;font-size:13px;padding:8px 16px;display:flex;gap:14px;position:sticky;top:0;z-index:20}}
.topbar a{{color:#e8edf5;text-decoration:none;font-weight:600}}
.doc{{max-width:820px;margin:0 auto;padding:28px 22px 60px;background:#fff}}
.badge{{letter-spacing:.16em;text-transform:uppercase;font-size:11px;color:var(--gold);font-weight:800}}
h1{{font-size:29px;color:var(--navy);margin:6px 0 4px;line-height:1.2}}
.meta{{color:var(--muted);font-size:14px;font-style:italic;margin-bottom:4px}}
.sub{{color:#33404f;font-size:16px;margin:6px 0 18px}}
.toc{{background:#f4f7fb;border:1px solid var(--line);border-radius:10px;padding:12px 16px 12px 34px;font-size:13.5px}}
.toc a{{color:var(--navy);text-decoration:none}}.toc li{{margin:3px 0}}
h2{{font-size:21px;color:var(--navy);border-bottom:2px solid var(--gold);padding-bottom:5px;margin:26px 0 10px}}
h3{{font-size:16px;color:var(--navy)}}
p{{margin:0 0 12px;text-align:justify}}
.hl{{background:#fff8ec;border-left:4px solid var(--gold);border-radius:8px;padding:8px 18px;margin:18px 0}}
.hl ul{{margin:6px 0}}
.gal{{display:grid;grid-template-columns:repeat(auto-fill,minmax(220px,1fr));gap:14px}}
figure{{margin:0}}figure img{{width:100%;height:170px;object-fit:cover;border-radius:8px;background:#eef2f8;border:1px solid var(--line)}}
figcaption{{font-size:12px;color:var(--muted);margin-top:4px}}
.refs a{{color:var(--navy)}}.refs li{{margin:4px 0;font-size:13.5px}}
.src{{margin-top:24px;padding-top:12px;border-top:1px dashed var(--line);color:var(--muted);font-size:12.5px}}
.foot{{margin-top:30px;padding-top:12px;border-top:1px solid var(--line);color:var(--muted);font-size:12px}}
.ttsbar{{position:fixed;left:0;right:0;bottom:0;background:#fff;border-top:1px solid var(--line);box-shadow:0 -4px 16px rgba(0,0,0,.08);display:flex;align-items:center;gap:10px;justify-content:center;padding:9px 12px;z-index:30;flex-wrap:wrap}}
.tb{{background:var(--navy);color:#fff;border:none;border-radius:999px;padding:9px 16px;font-size:14px;font-weight:700;cursor:pointer}}
.tb#ttsStop{{background:#eef2f8;color:var(--navy)}}
.trate{{font-size:12.5px;color:var(--muted)}}.trate input{{vertical-align:middle;width:90px}}
@media print{{.doc{{max-width:none}}a{{color:var(--ink)}}.no-print,.topbar{{display:none!important}}body{{padding-bottom:0}}}}
</style></head><body>
<div class="topbar no-print"><a href="../index.html">← Tất cả tài liệu</a><a href="../../trung-tam.html">🏠 Trang chủ</a></div>
<div class="doc">
<div class="badge">Cẩm nang Du lịch · Tài liệu thuyết minh nội bộ</div>
<h1>{esc(d.get('name_vi',''))}</h1>
<div class="meta">{esc(meta)}</div>
<div class="sub">{esc(d.get('subtitle',''))}</div>
<div class="toc"><b>MỤC LỤC</b><ol>{toc}</ol></div>
{hlbox}{body}{imgbox}{refbox}{srcbox}
<div class="foot">Tài liệu nội bộ — không dùng cho mục đích public. Nội dung biên soạn nguyên gốc, có dẫn nguồn dữ kiện. Cập nhật {TODAY}. · Cẩm nang Du lịch · lopmaybay@gmail.com</div>
</div>
{TTS_JS}
</body></html>"""


def build_docx(d, path):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    d = _fix_lists(normd(d))
    NAVY = RGBColor(0x15, 0x2c, 0x4e); GOLD = RGBColor(0xa9, 0x86, 0x3a); MUT = RGBColor(0x5c, 0x67, 0x73)
    LINK = RGBColor(0x20, 0x4a, 0x86)
    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Arial"; st.font.size = Pt(12)
    b = doc.add_paragraph(); r = b.add_run("CẨM NANG DU LỊCH NGA · TÀI LIỆU THUYẾT MINH NỘI BỘ")
    r.font.size = Pt(9); r.font.bold = True; r.font.color.rgb = GOLD
    t = doc.add_paragraph(); r = t.add_run(d.get("name_vi", "")); r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = NAVY
    meta = " · ".join([x for x in [d.get("name_ru"), d.get("name_en")] if x])
    if meta:
        p = doc.add_paragraph(); r = p.add_run(meta); r.italic = True; r.font.size = Pt(11); r.font.color.rgb = MUT
    if d.get("subtitle"):
        p = doc.add_paragraph(); r = p.add_run(d["subtitle"]); r.font.size = Pt(12.5)
    if d.get("highlights"):
        h = doc.add_paragraph(); r = h.add_run("★ Điểm nhấn"); r.bold = True; r.font.color.rgb = NAVY; r.font.size = Pt(13)
        for x in d["highlights"]:
            doc.add_paragraph(x, style="List Bullet")
    for s in d.get("sections", []):
        hd = doc.add_heading(level=1); r = hd.add_run(s.get("heading", "")); r.font.color.rgb = NAVY
        for p in s.get("paras", []):
            par = doc.add_paragraph(p); par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if d.get("images"):
        hd = doc.add_heading(level=1); r = hd.add_run("Hình ảnh (link chèn cho bản HTML)"); r.font.color.rgb = NAVY
        for im in d["images"]:
            p = doc.add_paragraph(); r = p.add_run("• " + im.get("caption", "") + " — "); r.font.size = Pt(11)
            r2 = p.add_run(im.get("url", "")); r2.font.size = Pt(10); r2.font.color.rgb = LINK
    if d.get("references"):
        hd = doc.add_heading(level=1); r = hd.add_run("Nguồn tham khảo & liên kết"); r.font.color.rgb = NAVY
        for rf in d["references"]:
            p = doc.add_paragraph(); r = p.add_run("• " + rf.get("title", "") + " — "); r.font.size = Pt(11)
            r2 = p.add_run(rf.get("url", "")); r2.font.size = Pt(10); r2.font.color.rgb = LINK
    if d.get("sources"):
        hd = doc.add_heading(level=2); r = hd.add_run("Nguồn dữ kiện"); r.font.color.rgb = MUT
        for x in d["sources"]:
            p = doc.add_paragraph(x); p.runs[0].font.size = Pt(10); p.runs[0].font.color.rgb = MUT
    f = doc.add_paragraph(); r = f.add_run("Tài liệu nội bộ — không dùng cho mục đích public. Biên soạn nguyên gốc, có dẫn nguồn. Cập nhật " + TODAY + ".")
    r.italic = True; r.font.size = Pt(9); r.font.color.rgb = MUT
    doc.save(path)


def main():
    files = sorted(glob.glob(os.path.join(IN, "doc_*.json")))
    if not files:
        print("Không có file doc_*.json trong _incoming."); return
    n = 0
    for fp in files:
        d = json.load(open(fp, encoding="utf-8"))
        reg = region_of(d)
        outdir = os.path.join(OUTROOT, reg); os.makedirs(outdir, exist_ok=True)
        slug = d.get("slug", os.path.basename(fp)[4:-5])
        with open(os.path.join(outdir, slug + ".html"), "w", encoding="utf-8") as f:
            f.write(build_html(d))
        try:
            build_docx(d, os.path.join(outdir, slug + ".docx"))
        except Exception as e:
            print("  (docx lỗi cho", slug, ":", e, ")")
        print(f"+ {reg}/{slug}: HTML+DOCX")
        n += 1
    print(f"Xong {n} tài liệu → {OUTROOT}/<vùng>/")
    # tự cập nhật trang index
    try:
        import subprocess
        subprocess.run(["python3", os.path.join(HERE, "build_docs_index.py")], check=False)
    except Exception:
        pass


if __name__ == "__main__":
    main()
